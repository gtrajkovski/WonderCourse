"""DOCX Textbook Exporter for Course Builder Studio.

Exports textbook chapters to Microsoft Word format with proper heading hierarchy,
glossary, references, and image placeholders.
"""

from pathlib import Path
from typing import Optional, List, Dict, Any

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.exporters.base_exporter import BaseExporter
from src.core.models import Course, TextbookChapter


class DOCXTextbookExporter(BaseExporter):
    """Exports course textbook chapters to Microsoft Word DOCX format.

    Creates a formatted Word document with:
    - Title page with course title and description
    - Chapters as Heading 1
    - Sections as Heading 2
    - Aggregated glossary across all chapters
    - Aggregated references across all chapters
    - Image placeholders as italic text
    """

    @property
    def format_name(self) -> str:
        """Human-readable name of the export format."""
        return "Microsoft Word"

    @property
    def file_extension(self) -> str:
        """File extension for exported files."""
        return ".docx"

    def export(self, course: Course, filename: Optional[str] = None) -> Path:
        """Export textbook chapters to DOCX file.

        Args:
            course: Course object containing textbook_chapters.
            filename: Optional filename (without extension). If None, uses course title.

        Returns:
            Path to the exported DOCX file.
        """
        output_path = self.get_output_path(course, filename)

        # Create new document
        doc = Document()

        # Add title page
        self._add_title_page(doc, course)

        # Collect all glossary terms and references for aggregation
        all_glossary_terms: List[Dict[str, str]] = []
        all_references: List[Dict[str, str]] = []

        # Add each chapter
        for chapter_num, chapter in enumerate(course.textbook_chapters, start=1):
            self._add_chapter(doc, chapter, chapter_num)

            # Collect glossary and references
            all_glossary_terms.extend(chapter.glossary_terms)
            all_references.extend(chapter.references)

        # Add aggregated glossary section (if any terms exist)
        if all_glossary_terms:
            self._add_glossary(doc, all_glossary_terms)

        # Add aggregated references section (if any references exist)
        if all_references:
            self._add_references(doc, all_references)

        # Save document
        doc.save(str(output_path))

        return output_path

    def _add_title_page(self, doc: Document, course: Course) -> None:
        """Add title page with course information.

        Args:
            doc: Word document to add title page to.
            course: Course object for title and metadata.
        """
        # Add title
        title_para = doc.add_paragraph(course.title)
        title_para.style = "Title"
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add description if present
        if course.description:
            desc_para = doc.add_paragraph(course.description)
            desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a page break after title page (only if there are chapters)
        if course.textbook_chapters:
            doc.add_page_break()

    def _add_chapter(self, doc: Document, chapter: TextbookChapter, chapter_num: int) -> None:
        """Add a chapter to the document.

        Args:
            doc: Word document to add chapter to.
            chapter: TextbookChapter object with chapter content.
            chapter_num: Chapter number for ordering.
        """
        # Chapter title as Heading 1
        chapter_heading = f"Chapter {chapter_num}: {chapter.title}"
        doc.add_heading(chapter_heading, level=1)

        # Add each section
        for section in chapter.sections:
            self._add_section(doc, section)

        # Add image placeholders for this chapter
        if chapter.image_placeholders:
            self._add_image_placeholders(doc, chapter.image_placeholders)

    def _add_section(self, doc: Document, section: Dict[str, str]) -> None:
        """Add a section to the document.

        Args:
            doc: Word document to add section to.
            section: Section dictionary with 'heading' and 'content' keys.
        """
        # Section heading as Heading 2
        heading = section.get("heading", "")
        if heading:
            doc.add_heading(heading, level=2)

        # Section content as normal paragraph
        content = section.get("content", "")
        if content:
            doc.add_paragraph(content)

    def _add_image_placeholders(self, doc: Document, placeholders: List[Dict[str, str]]) -> None:
        """Add image placeholders as italic text.

        Args:
            doc: Word document to add placeholders to.
            placeholders: List of image placeholder dictionaries.
        """
        for placeholder in placeholders:
            figure_num = placeholder.get("figure_number", "Figure")
            caption = placeholder.get("caption", "")
            alt_text = placeholder.get("alt_text", "")

            # Create placeholder text
            placeholder_text = f"[{figure_num}: {caption}]"
            if alt_text:
                placeholder_text += f" ({alt_text})"

            # Add as italic paragraph
            para = doc.add_paragraph()
            run = para.add_run(placeholder_text)
            run.italic = True

    def _add_glossary(self, doc: Document, glossary_terms: List[Dict[str, str]]) -> None:
        """Add aggregated glossary section.

        Args:
            doc: Word document to add glossary to.
            glossary_terms: List of all glossary term dictionaries.
        """
        # Add page break before glossary
        doc.add_page_break()

        # Glossary heading
        doc.add_heading("Glossary", level=1)

        # Deduplicate and sort terms
        seen_terms = set()
        unique_terms = []
        for term_dict in glossary_terms:
            term = term_dict.get("term", "")
            if term and term not in seen_terms:
                seen_terms.add(term)
                unique_terms.append(term_dict)

        # Sort alphabetically
        unique_terms.sort(key=lambda x: x.get("term", "").lower())

        # Add each term
        for term_dict in unique_terms:
            term = term_dict.get("term", "")
            definition = term_dict.get("definition", "")

            para = doc.add_paragraph()
            # Term in bold
            term_run = para.add_run(f"{term}: ")
            term_run.bold = True
            # Definition in normal text
            para.add_run(definition)

    def _add_references(self, doc: Document, references: List[Dict[str, str]]) -> None:
        """Add aggregated references section.

        Args:
            doc: Word document to add references to.
            references: List of all reference dictionaries.
        """
        # Add page break before references
        doc.add_page_break()

        # References heading
        doc.add_heading("References", level=1)

        # Deduplicate references by citation text
        seen_citations = set()
        unique_refs = []
        for ref_dict in references:
            citation = ref_dict.get("citation", "")
            if citation and citation not in seen_citations:
                seen_citations.add(citation)
                unique_refs.append(ref_dict)

        # Sort alphabetically by citation
        unique_refs.sort(key=lambda x: x.get("citation", "").lower())

        # Add each reference
        for ref_dict in unique_refs:
            citation = ref_dict.get("citation", "")
            url = ref_dict.get("url", "")

            para = doc.add_paragraph(citation, style="List Bullet")

            # Add URL if present
            if url:
                url_para = doc.add_paragraph()
                url_para.paragraph_format.left_indent = Inches(0.5)
                url_run = url_para.add_run(url)
                url_run.italic = True
