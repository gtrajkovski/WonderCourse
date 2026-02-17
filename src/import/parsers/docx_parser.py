"""DOCX document parser for content import.

Uses python-docx for structured extraction and mammoth for semantic HTML conversion.
"""

from datetime import datetime
from typing import Union
import io

try:
    from docx import Document
    import mammoth
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .base_parser import BaseParser, ParseResult


class DOCXParser(BaseParser):
    """Parser for Microsoft Word .docx files.

    Extracts structured content including headings, paragraphs, tables, and lists.
    Detects content type based on document structure.
    """

    def can_parse(self, source: Union[str, bytes], filename: str = None) -> bool:
        """Detect if source is a DOCX file.

        Args:
            source: File bytes or path
            filename: Optional filename for extension checking

        Returns:
            True if source appears to be a DOCX file
        """
        if not DOCX_AVAILABLE:
            return False

        # Check filename extension
        if filename and filename.lower().endswith('.docx'):
            return True

        # Check DOCX magic bytes (ZIP header: PK)
        if isinstance(source, bytes) and len(source) >= 4:
            return source[:2] == b'PK'

        return False

    def parse(self, source: Union[str, bytes], filename: str = None) -> ParseResult:
        """Parse DOCX content into structured format.

        Args:
            source: DOCX file bytes
            filename: Optional filename for provenance

        Returns:
            ParseResult with extracted content

        Raises:
            ValueError: If source cannot be parsed as DOCX
        """
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx and mammoth libraries not installed")

        if not self.can_parse(source, filename):
            raise ValueError("Source is not a valid DOCX file")

        # Convert bytes to file-like object if needed
        if isinstance(source, bytes):
            file_obj = io.BytesIO(source)
        else:
            raise ValueError("Source must be bytes for DOCX parsing")

        # Extract structured content with python-docx
        doc = Document(file_obj)

        paragraphs = []
        tables = []
        warnings = []

        # Extract paragraphs with styles
        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            para_data = {
                'text': para.text,
                'style': para.style.name if para.style else 'Normal',
                'runs': []
            }

            for run in para.runs:
                para_data['runs'].append({
                    'text': run.text,
                    'bold': run.bold or False,
                    'italic': run.italic or False,
                    'underline': run.underline or False
                })

            paragraphs.append(para_data)

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(cells)
            tables.append(rows)

        # Extract metadata from document properties
        metadata = {
            'author': doc.core_properties.author or 'Unknown',
            'title': doc.core_properties.title or '',
            'created_date': str(doc.core_properties.created) if doc.core_properties.created else None,
            'paragraph_count': len(paragraphs),
            'table_count': len(tables),
            'format': 'docx'
        }

        # Convert to HTML with mammoth for semantic structure
        file_obj.seek(0)  # Reset to beginning
        mammoth_result = mammoth.convert_to_html(file_obj)
        html_content = mammoth_result.value

        # Add mammoth warnings
        for message in mammoth_result.messages:
            if hasattr(message, 'message'):
                warnings.append(f"Conversion: {message.message}")

        # Detect content type from structure
        content_type = self._detect_content_type(paragraphs)

        # Calculate word count
        word_count = sum(len(p['text'].split()) for p in paragraphs)
        metadata['word_count'] = word_count

        content = {
            'paragraphs': paragraphs,
            'tables': tables,
            'html': html_content,
            'title': metadata['title'] or (paragraphs[0]['text'][:100] if paragraphs else '')
        }

        provenance = {
            'filename': filename or 'unknown.docx',
            'import_time': datetime.now().isoformat(),
            'original_format': 'docx',
            'parser': 'DOCXParser'
        }

        return ParseResult(
            content_type=content_type,
            content=content,
            metadata=metadata,
            warnings=warnings,
            provenance=provenance
        )

    def _detect_content_type(self, paragraphs):
        """Detect content type from document structure.

        Args:
            paragraphs: List of paragraph dictionaries

        Returns:
            Detected content type string
        """
        if not paragraphs:
            return 'reading'

        # Count headings
        heading_count = sum(1 for p in paragraphs if 'Heading' in p['style'])

        # Check for numbered lists (step-by-step instructions)
        numbered_lines = sum(1 for p in paragraphs
                            if p['text'].strip() and p['text'].strip()[0].isdigit())

        # Heuristics for content type detection
        if numbered_lines > len(paragraphs) * 0.3:
            return 'lab'  # Step-by-step instructions
        elif heading_count > len(paragraphs) * 0.2:
            return 'reading'  # Structured document with sections
        else:
            return 'reading'  # Default to reading material
