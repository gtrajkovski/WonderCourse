"""Tests for DOCXTextbookExporter - TDD tests for DOCX textbook export."""

import pytest
from pathlib import Path
from unittest.mock import Mock, patch
from docx import Document

from src.exporters.docx_textbook import DOCXTextbookExporter
from src.exporters.base_exporter import BaseExporter
from src.core.models import Course, TextbookChapter


@pytest.fixture
def temp_output_dir(tmp_path):
    """Provide a temporary directory for output files."""
    return tmp_path


@pytest.fixture
def exporter(temp_output_dir):
    """Create DOCXTextbookExporter instance with temp output directory."""
    return DOCXTextbookExporter(output_dir=temp_output_dir)


@pytest.fixture
def course_with_textbook():
    """Create a course with textbook chapters."""
    course = Course(
        id="course_test",
        title="Introduction to Machine Learning",
        description="A comprehensive course on ML fundamentals",
    )

    # Chapter 1 with sections, glossary, and references
    chapter1 = TextbookChapter(
        id="ch_001",
        title="Neural Networks Fundamentals",
        sections=[
            {"heading": "Introduction", "content": "Neural networks are computing systems inspired by biological neural networks."},
            {"heading": "Architecture", "content": "A typical neural network consists of input, hidden, and output layers."},
            {"heading": "Training", "content": "Networks learn through a process called backpropagation."}
        ],
        glossary_terms=[
            {"term": "Neural Network", "definition": "A computing system inspired by biological neural networks."},
            {"term": "Backpropagation", "definition": "An algorithm for training neural networks."}
        ],
        references=[
            {"citation": "LeCun, Y. (2015). Deep learning. Nature.", "url": "https://nature.com/articles"},
            {"citation": "Goodfellow, I. (2016). Deep Learning. MIT Press.", "url": ""}
        ],
        image_placeholders=[
            {"figure_number": "Figure 1.1", "caption": "Neural network architecture", "alt_text": "Diagram of layers"}
        ],
        word_count=500
    )

    # Chapter 2 - minimal
    chapter2 = TextbookChapter(
        id="ch_002",
        title="Deep Learning Applications",
        sections=[
            {"heading": "Computer Vision", "content": "Deep learning has transformed computer vision tasks."}
        ],
        glossary_terms=[],
        references=[],
        word_count=100
    )

    course.textbook_chapters = [chapter1, chapter2]
    return course


@pytest.fixture
def course_without_textbook():
    """Create a course without textbook chapters."""
    return Course(
        id="course_empty",
        title="Empty Textbook Course",
        description="A course with no textbook content",
        textbook_chapters=[]
    )


class TestDOCXTextbookExporterInheritance:
    """Tests for BaseExporter inheritance."""

    def test_inherits_from_base_exporter(self, exporter):
        """DOCXTextbookExporter should inherit from BaseExporter."""
        assert isinstance(exporter, BaseExporter)

    def test_format_name_property(self, exporter):
        """format_name should return 'Microsoft Word'."""
        assert exporter.format_name == "Microsoft Word"

    def test_file_extension_property(self, exporter):
        """file_extension should return '.docx'."""
        assert exporter.file_extension == ".docx"


class TestDOCXTextbookExport:
    """Tests for basic DOCX export functionality."""

    def test_export_creates_docx_file(self, exporter, course_with_textbook):
        """Export should create a .docx file."""
        output_path = exporter.export(course_with_textbook)

        assert output_path.exists()
        assert output_path.suffix == ".docx"

    def test_export_uses_custom_filename(self, exporter, course_with_textbook):
        """Export should use custom filename when provided."""
        output_path = exporter.export(course_with_textbook, filename="custom_textbook")

        assert output_path.name == "custom_textbook.docx"
        assert output_path.exists()

    def test_export_uses_course_title_as_default_filename(self, exporter, course_with_textbook):
        """Export should use sanitized course title as default filename."""
        output_path = exporter.export(course_with_textbook)

        # Course title "Introduction to Machine Learning" should become valid filename
        assert "Introduction" in output_path.stem or "Machine" in output_path.stem

    def test_export_returns_path_object(self, exporter, course_with_textbook):
        """Export should return a Path object."""
        output_path = exporter.export(course_with_textbook)

        assert isinstance(output_path, Path)


class TestDOCXTextbookContent:
    """Tests for DOCX document content structure."""

    def test_document_has_title_page(self, exporter, course_with_textbook):
        """Exported document should have title page with course title."""
        output_path = exporter.export(course_with_textbook)
        doc = Document(output_path)

        # Title should appear as first heading (Title style)
        first_para = doc.paragraphs[0]
        assert course_with_textbook.title in first_para.text

    def test_document_has_chapters_as_heading1(self, exporter, course_with_textbook):
        """Chapters should use Heading 1 style."""
        output_path = exporter.export(course_with_textbook)
        doc = Document(output_path)

        # Find all Heading 1 paragraphs
        heading1_paragraphs = [
            p for p in doc.paragraphs
            if p.style and p.style.name == "Heading 1"
        ]

        # Should have at least 2 chapter headings
        assert len(heading1_paragraphs) >= 2

        # Verify chapter titles appear in headings
        heading_texts = [p.text for p in heading1_paragraphs]
        assert any("Neural Networks" in h for h in heading_texts)

    def test_document_has_sections_as_heading2(self, exporter, course_with_textbook):
        """Sections should use Heading 2 style."""
        output_path = exporter.export(course_with_textbook)
        doc = Document(output_path)

        # Find all Heading 2 paragraphs
        heading2_paragraphs = [
            p for p in doc.paragraphs
            if p.style and p.style.name == "Heading 2"
        ]

        # Should have section headings from chapter 1 (Introduction, Architecture, Training)
        assert len(heading2_paragraphs) >= 3

        heading_texts = [p.text for p in heading2_paragraphs]
        assert any("Introduction" in h for h in heading_texts)
        assert any("Architecture" in h for h in heading_texts)

    def test_document_contains_section_content(self, exporter, course_with_textbook):
        """Section content should appear in the document."""
        output_path = exporter.export(course_with_textbook)
        doc = Document(output_path)

        full_text = "\n".join([p.text for p in doc.paragraphs])

        # Content from sections should appear
        assert "computing systems inspired by biological" in full_text
        assert "backpropagation" in full_text


class TestDOCXTextbookGlossary:
    """Tests for glossary section in DOCX."""

    def test_document_has_glossary_section(self, exporter, course_with_textbook):
        """Document should have a Glossary section."""
        output_path = exporter.export(course_with_textbook)
        doc = Document(output_path)

        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Glossary" in full_text

    def test_glossary_contains_terms(self, exporter, course_with_textbook):
        """Glossary should contain terms and definitions."""
        output_path = exporter.export(course_with_textbook)
        doc = Document(output_path)

        full_text = "\n".join([p.text for p in doc.paragraphs])

        # Glossary terms from chapter 1
        assert "Neural Network" in full_text
        assert "Backpropagation" in full_text
        # Definitions should also appear
        assert "computing system inspired by biological" in full_text


class TestDOCXTextbookReferences:
    """Tests for references section in DOCX."""

    def test_document_has_references_section(self, exporter, course_with_textbook):
        """Document should have a References section."""
        output_path = exporter.export(course_with_textbook)
        doc = Document(output_path)

        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "References" in full_text

    def test_references_contain_citations(self, exporter, course_with_textbook):
        """References section should contain citations."""
        output_path = exporter.export(course_with_textbook)
        doc = Document(output_path)

        full_text = "\n".join([p.text for p in doc.paragraphs])

        # Citations from chapter 1
        assert "LeCun" in full_text
        assert "Goodfellow" in full_text


class TestDOCXTextbookImagePlaceholders:
    """Tests for image placeholder handling."""

    def test_image_placeholders_as_italic_text(self, exporter, course_with_textbook):
        """Image placeholders should appear as italic text."""
        output_path = exporter.export(course_with_textbook)
        doc = Document(output_path)

        full_text = "\n".join([p.text for p in doc.paragraphs])

        # Image placeholder info should appear
        assert "Figure 1.1" in full_text
        assert "Neural network architecture" in full_text


class TestDOCXTextbookEdgeCases:
    """Tests for edge cases and error handling."""

    def test_export_empty_textbook_creates_minimal_doc(self, exporter, course_without_textbook):
        """Exporting course without textbook chapters should create minimal doc."""
        output_path = exporter.export(course_without_textbook)

        assert output_path.exists()
        doc = Document(output_path)

        # Should at least have the title
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert course_without_textbook.title in full_text

    def test_export_chapter_without_glossary(self, exporter):
        """Chapter without glossary terms should export without errors."""
        course = Course(
            id="course_test",
            title="Minimal Course",
            textbook_chapters=[
                TextbookChapter(
                    id="ch_001",
                    title="Chapter One",
                    sections=[{"heading": "Section", "content": "Content here."}],
                    glossary_terms=[],
                    references=[]
                )
            ]
        )

        output_path = exporter.export(course)
        assert output_path.exists()

    def test_export_chapter_without_references(self, exporter):
        """Chapter without references should export without errors."""
        course = Course(
            id="course_test",
            title="No References Course",
            textbook_chapters=[
                TextbookChapter(
                    id="ch_001",
                    title="Chapter Without Refs",
                    sections=[{"heading": "Section", "content": "No references here."}],
                    glossary_terms=[{"term": "Test", "definition": "A test term."}],
                    references=[]
                )
            ]
        )

        output_path = exporter.export(course)
        assert output_path.exists()

    def test_special_characters_in_title(self, exporter):
        """Course with special characters in title should export safely."""
        course = Course(
            id="course_special",
            title="Course: Advanced C++ & Python/JavaScript",
            textbook_chapters=[
                TextbookChapter(
                    id="ch_001",
                    title="Chapter 1",
                    sections=[{"heading": "Test", "content": "Test content."}]
                )
            ]
        )

        output_path = exporter.export(course)
        assert output_path.exists()
        # Filename should not contain problematic characters
        assert ":" not in output_path.name
        assert "/" not in output_path.name
